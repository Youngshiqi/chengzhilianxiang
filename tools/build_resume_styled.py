# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"C:\Users\yangshiqi\Desktop\2300130314杨诗琪_样式优化版.docx"

FONT = "Microsoft YaHei"
NAVY = "183B56"
TEAL = "0F766E"
BLUE = "2563EB"
INK = "1F2937"
MUTED = "64748B"
LINE = "D6E0EA"
PALE = "F4F8FB"
PALE_TEAL = "E8F5F3"
PALE_BLUE = "EEF5FF"
WHITE = "FFFFFF"


def set_run(run, size=9.5, bold=False, color=INK, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_run(paragraph, text, size=9.5, bold=False, color=INK):
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return run


def set_para(paragraph, before=0, after=2, line=1.0, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def add_section_heading(doc, title):
    p = doc.add_paragraph()
    set_para(p, before=5, after=2, line=1.0)
    add_run(p, title, size=10.6, bold=True, color=NAVY)
    p_pr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "7")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), TEAL)
    bdr.append(bottom)
    p_pr.append(bdr)


def add_bullet(doc, text, highlight=None):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, after=1, line=1.03)
    if highlight and text.startswith(highlight):
        add_run(p, highlight, size=8.8, bold=True, color=NAVY)
        add_run(p, text[len(highlight):], size=8.8, color=INK)
    else:
        add_run(p, text, size=8.8, color=INK)
    return p


def add_tag_row(doc, tags):
    table = doc.add_table(rows=1, cols=len(tags))
    table.autofit = False
    set_table_width(table, 9360)
    for i, tag in enumerate(tags):
        cell = table.cell(0, i)
        shade_cell(cell, PALE_TEAL if i % 2 == 0 else PALE_BLUE)
        set_cell_border(cell, "E2E8F0", "3")
        set_cell_margins(cell, top=55, bottom=55, start=70, end=70)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_para(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p, tag, size=8.1, bold=True, color=NAVY)


def add_skill_grid(doc):
    rows = [
        ("后端", "Python / FastAPI / Pydantic / SQLAlchemy Async / RESTful API / JWT"),
        ("数据", "MySQL / Redis 缓存、Geo、分布式锁 / MongoDB / Elasticsearch"),
        ("AI", "LangChain + 百炼 LLM / Prompt 设计 / 结构化输出 / 多模态视觉验收 / AI日志沉淀"),
        ("工程", "Vue 3 / Element Plus / ECharts / Jinja2 / Docker Compose / Git / 测试脚本"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    set_table_width(table, 9360)
    for i, (label, value) in enumerate(rows):
        label_cell = table.cell(i, 0)
        value_cell = table.cell(i, 1)
        label_cell.width = Cm(2.2)
        value_cell.width = Cm(14.3)
        shade_cell(label_cell, TEAL if i % 2 == 0 else NAVY)
        shade_cell(value_cell, "FBFCFE")
        for cell in (label_cell, value_cell):
            set_cell_border(cell, "D8E3ED", "4")
            set_cell_margins(cell, top=65, bottom=65, start=100, end=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = label_cell.paragraphs[0]
        set_para(p0, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(p0, label, size=8.8, bold=True, color=WHITE)
        p1 = value_cell.paragraphs[0]
        set_para(p1, after=0, line=1.0)
        add_run(p1, value, size=8.7, color=INK)


def add_project(doc, title, date, stack, bullets):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table, 9360)
    cell = table.cell(0, 0)
    shade_cell(cell, PALE)
    set_cell_border(cell, "D9E5EF", "4")
    set_cell_margins(cell, top=70, bottom=70, start=120, end=120)
    p = cell.paragraphs[0]
    set_para(p, after=0, line=1.0)
    add_run(p, title, size=9.7, bold=True, color=NAVY)
    add_run(p, f"  |  {date}", size=8.3, bold=True, color=TEAL)
    p2 = cell.add_paragraph()
    set_para(p2, before=1, after=0, line=1.0)
    add_run(p2, stack, size=8.1, color=MUTED)
    for item in bullets:
        add_bullet(doc, item)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(0.95)
    section.bottom_margin = Cm(0.85)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.03

    bullet = styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    bullet.font.size = Pt(8.8)
    bullet.paragraph_format.left_indent = Cm(0.38)
    bullet.paragraph_format.first_line_indent = Cm(-0.18)
    bullet.paragraph_format.space_after = Pt(1)
    bullet.paragraph_format.line_spacing = 1.03

    # Top identity band.
    top = doc.add_table(rows=1, cols=2)
    top.autofit = False
    set_table_width(top, 9360)
    top.cell(0, 0).width = Cm(9.5)
    top.cell(0, 1).width = Cm(7.0)
    for cell in top.rows[0].cells:
        shade_cell(cell, NAVY)
        set_cell_border(cell, NAVY, "1")
        set_cell_margins(cell, top=145, bottom=135, start=150, end=150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left = top.cell(0, 0).paragraphs[0]
    set_para(left, after=0, line=1.0)
    add_run(left, "杨诗琪", size=20.5, bold=True, color=WHITE)
    add_run(left, "  AI应用开发 / Python后端开发", size=10.2, bold=True, color="BFE7E3")
    sub = top.cell(0, 0).add_paragraph()
    set_para(sub, before=2, after=0, line=1.0)
    add_run(sub, "FastAPI 异步后端 · LangChain AI 工作流 · 多数据源系统设计", size=8.7, color="D8EAF5")

    right = top.cell(0, 1).paragraphs[0]
    set_para(right, after=0, line=1.18, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(right, "电话：130xxxxxxxx\n", size=8.5, color=WHITE)
    add_run(right, "邮箱：xxxxxx@xxx.com\n", size=8.5, color=WHITE)
    add_run(right, "本科在读 / 2027届 / 居住地：XXXX", size=8.5, color=WHITE)

    p = doc.add_paragraph()
    set_para(p, before=4, after=3, line=1.08)
    add_run(
        p,
        "计算机科学与技术本科在读，主攻 Python 后端与 AI 应用工程化。具备 FastAPI 异步服务、"
        "SQLAlchemy/MySQL、Redis、MongoDB、Elasticsearch、RabbitMQ、Dify 工作流及 Vue 前端协作经验；"
        "能独立完成从需求分析、数据建模、接口开发到联调测试的完整项目闭环。",
        size=8.9,
        color=INK,
    )

    add_tag_row(doc, ["FastAPI", "LangChain", "Redis Geo", "RabbitMQ", "MySQL", "MongoDB", "Elasticsearch", "Vue 3"])

    add_section_heading(doc, "教育背景")
    edu = doc.add_paragraph()
    set_para(edu, after=1, line=1.0)
    add_run(edu, "软件学院 / 计算机科学与技术 / 本科", size=9.2, bold=True, color=NAVY)
    add_run(edu, "  |  2023.09 - 2027.06", size=8.6, color=MUTED)
    add_bullet(doc, "核心课程：C语言、Java程序设计、HTML5应用开发、数据结构、计算机网络、操作系统、数据库原理。")

    add_section_heading(doc, "专业技能")
    add_skill_grid(doc)

    add_section_heading(doc, "项目经历")
    add_project(
        doc,
        "城市公共设施智能报修与派单系统",
        "2026.06",
        "FastAPI / Vue 3 / MySQL / Redis / MongoDB / Elasticsearch / RabbitMQ / LangChain + 百炼",
        [
            "构建三段式 AI 工作流：基于 LangChain + 百炼 qwen-vl-max-latest，实现报修 NLP 解析、派单评分、多模态视觉验收三个结构化输出服务，LLM 不可用时自动降级为关键词匹配与确定性算法，保障业务不断。",
            "负责核心后端业务编排：实现市民报修→Redis热缓存→AI解析回写→MongoDB日志→ES同步消息→RabbitMQ超时派单的全链路，以 MySQL 作为同步落地点，其余增强能力异步容错，保证报修主流程稳定返回受理回执。",
            "设计四库分层数据模型：MySQL 承载用户/工单/结算等事务数据，Redis 承载工单状态缓存、Geo空间计算和分布式锁，MongoDB 承载维修记录/附件/审计/AI日志，Elasticsearch 支撑全文检索与统计聚合。",
            "实现智能派单流程：基于 Redis Geo 半径筛选候选维修员，引入高德驾车距离修正与5分钟缓存，使用 Redis SETNX 锁防止并发双派，并通过 RabbitMQ 延迟队列实现10分钟无人接单自动升级强制指派。",
            "实现 ES 可靠同步：基于 RabbitMQ 异步投递 + 指数退避重试 + 死信队列兜底，重试耗尽后路由至 DLQ 避免消息丢失，保障最终一致性。",
        ],
    )
    add_project(
        doc,
        "Blog Log AI System 博客日志AI分析系统",
        "2026.05",
        "FastAPI / SQLAlchemy Async / MySQL / Dify API / Jinja2 / JavaScript",
        [
            "独立搭建博客管理与 AI 分析平台，完成注册登录、JWT 鉴权、博客 CRUD、分页查询、标题唯一性校验、Markdown 内容管理和用户数据隔离。",
            "封装 Dify 工作流调用，支持单篇博客摘要/关键词/难度/分类分析、全局主题聚类分析、个性化学习路线推荐，并将原始响应与结构化结果落库，便于复查与二次展示。",
            "建设统计分析接口：实现总博客数、月度新增、总字数、今日字数、发布/草稿分布、TOP博客、写作时段等指标；编写 Markdown 清洗函数，提升字数统计准确性。",
            "按 API-Service-CRUD-Model 分层组织代码，配套统一日志、异常处理、响应结构和测试脚本，降低接口扩展与问题定位成本。",
        ],
    )

    add_section_heading(doc, "个人优势")
    add_bullet(doc, "能把 AI 能力落到具体业务闭环中，关注结构化输出、失败降级、日志追踪和数据沉淀。")
    add_bullet(doc, "熟悉异步后端与多数据源协作，项目覆盖缓存、消息队列、文档存储、全文检索和前后端联调。")
    add_bullet(doc, "文档意识较强，能维护 PRD、接口说明和测试脚本，适合 AI 应用开发、Python 后端开发、政企数字化系统开发等岗位。")

    add_section_heading(doc, "补充信息")
    tail = doc.add_paragraph()
    set_para(tail, after=0, line=1.05)
    add_run(tail, "求职方向：", size=8.7, bold=True, color=NAVY)
    add_run(tail, "AI应用开发工程师、Python后端开发工程师、后端开发实习生。", size=8.7)
    tail2 = doc.add_paragraph()
    set_para(tail2, after=0, line=1.05)
    add_run(tail2, "可面试展开：", size=8.7, bold=True, color=NAVY)
    add_run(tail2, "FastAPI 异步接口设计、Redis Geo 派单、RabbitMQ 延迟/死信队列、LangChain 结构化输出、ES 聚合统计。", size=8.7)

    doc.save(OUT)


if __name__ == "__main__":
    build()
