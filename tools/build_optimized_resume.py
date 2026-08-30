from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"C:\Users\yangshiqi\Desktop\2300130314杨诗琪_优化版.docx"


BLUE = "1F4E79"
LIGHT_BLUE = "EAF2F8"
GRAY = "666666"
DARK = "222222"


def set_run_font(run, size=10.5, bold=False, color=DARK, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(paragraph, before=0, after=3, line=1.0, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_text(paragraph, text, size=10.5, bold=False, color=DARK):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def add_section_heading(doc, title):
    p = doc.add_paragraph()
    set_para(p, before=7, after=2)
    add_text(p, title, size=12, bold=True, color=BLUE)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), BLUE)
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, after=2, line=1.08)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, size=9.4, bold=True)
        add_text(p, text[len(bold_prefix):], size=9.4)
    else:
        add_text(p, text, size=9.4)
    return p


def add_project(doc, title, date, stack, bullets):
    p = doc.add_paragraph()
    set_para(p, before=4, after=1, line=1.0)
    add_text(p, title, size=10.2, bold=True, color=DARK)
    add_text(p, "  |  " + stack, size=9.2, color=GRAY)
    if date:
        add_text(p, "  |  " + date, size=9.2, color=GRAY)
    for b in bullets:
        add_bullet(doc, b)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)

    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(9.4)
        st.paragraph_format.left_indent = Cm(0.48)
        st.paragraph_format.first_line_indent = Cm(-0.22)
        st.paragraph_format.space_after = Pt(2)
        st.paragraph_format.line_spacing = 1.08

    # Header block.
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.6)
    table.columns[1].width = Cm(7.3)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
            for p in cell.paragraphs:
                set_para(p, after=0)

    left = table.cell(0, 0).paragraphs[0]
    add_text(left, "杨诗琪", size=22, bold=True, color=BLUE)
    add_text(left, "  AI应用开发 / Python后端开发", size=11, color=GRAY)

    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(right, "电话：130xxxxxxxx  |  邮箱：xxxxxx@xxx.com\n", size=9.2, color=GRAY)
    add_text(right, "居住地：XXXX  |  本科在读  |  2027届", size=9.2, color=GRAY)

    p = doc.add_paragraph()
    set_para(p, before=4, after=4, line=1.15)
    add_text(
        p,
        "计算机科学与技术本科在读，主攻 Python 后端与 AI 应用工程化。具备 FastAPI 异步服务、"
        "SQLAlchemy/MySQL、Redis、MongoDB、Elasticsearch、RabbitMQ、Dify 工作流及 Vue 前端协作经验；"
        "能独立完成从需求分析、数据建模、接口开发到联调测试的完整项目闭环。",
        size=9.7,
    )

    add_section_heading(doc, "教育背景")
    edu = doc.add_paragraph()
    set_para(edu, after=1)
    add_text(edu, "软件学院 / 计算机科学与技术 / 本科", size=10, bold=True)
    add_text(edu, "  |  2023.09 - 2027.06", size=9.4, color=GRAY)
    add_bullet(doc, "核心课程：C语言、Java程序设计、HTML5应用开发、数据结构、计算机网络、操作系统、数据库原理。")

    add_section_heading(doc, "专业技能")
    skill_table = doc.add_table(rows=4, cols=2)
    skill_table.autofit = False
    widths = [Cm(3.0), Cm(13.8)]
    skills = [
        ("后端开发", "Python、FastAPI、Pydantic、SQLAlchemy 2.0 异步模式、RESTful API、JWT 鉴权、统一响应与异常处理"),
        ("数据与中间件", "MySQL、Redis 缓存/Geo/分布式锁、MongoDB 文档存储、Elasticsearch 检索聚合、RabbitMQ 异步消息"),
        ("AI 应用", "Dify 工作流接入、提示词设计、AI 结果结构化解析、内容分析/分类/路线推荐、AI 验收与日志沉淀"),
        ("前端与工程", "Vue 3、Element Plus、ECharts、Jinja2、原生 JavaScript、Docker Compose、Git、接口联调与测试脚本"),
    ]
    for i, (label, text) in enumerate(skills):
        row = skill_table.rows[i]
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        shade_cell(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = row.cells[0].paragraphs[0]
        set_para(p0, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p0, label, size=9.2, bold=True, color=BLUE)
        p1 = row.cells[1].paragraphs[0]
        set_para(p1, after=0, line=1.05)
        add_text(p1, text, size=9.0)

    add_section_heading(doc, "项目经历")
    add_project(
        doc,
        "城市公共设施智能报修与派单系统",
        "2026.06",
        "FastAPI / Vue 3 / MySQL / Redis / MongoDB / Elasticsearch / RabbitMQ / Dify",
        [
            "负责核心后端业务编排：实现市民报修、工单热缓存、图片元数据、AI解析日志、ES同步消息、超时派单检查等链路，采用 MySQL 作为同步落地点，其余增强能力异步容错，保证报修主流程可稳定返回受理回执。",
            "设计四库分层数据模型：MySQL 承载用户/工单/结算等事务数据，Redis 承载工单状态缓存、Geo空间计算和分布式锁，MongoDB 承载维修记录/附件/审计/AI日志，Elasticsearch 支撑全文检索与统计聚合。",
            "实现智能派单流程：基于 Redis Geo 进行候选维修员半径筛选，引入高德驾车距离修正与5分钟缓存，使用 Redis SETNX 锁防止并发双派，并通过 RabbitMQ 延迟消息处理无人接单兜底。",
            "参与三端页面联调：市民端报修/进度，维修员H5接单/完工/绩效，管理后台工单、调度、设施、结算、驾驶舱等模块，配合 ECharts 展示运营指标。",
        ],
    )
    add_project(
        doc,
        "Blog Log AI System 博客日志AI分析系统",
        "2026.05",
        "FastAPI / SQLAlchemy Async / MySQL / Dify API / Jinja2 / JavaScript",
        [
            "独立搭建博客管理与 AI 分析平台，完成注册登录、JWT 鉴权、博客 CRUD、分页查询、标题唯一性校验、Markdown 内容管理和用户数据隔离。",
            "封装 Dify 工作流调用，支持单篇博客摘要/关键词/难度/分类分析、全局主题聚类分析、个性化学习路线推荐，并将原始响应与结构化结果落库，便于复查与二次展示。",
            "建设统计分析接口：实现总博客数、月度新增、总字数、今日字数、发布/草稿分布、TOP博客、写作时段等指标；编写 Markdown 清洗函数，提升字数统计准确性。",
            "按 API-Service-CRUD-Model 分层组织代码，配套统一日志、异常处理、响应结构和测试脚本，降低接口扩展与问题定位成本。",
        ],
    )

    add_section_heading(doc, "个人优势")
    for item in [
        "能把 AI 能力落到具体业务闭环中，不停留在调用模型：关注结构化输出、失败降级、日志追踪和数据沉淀。",
        "熟悉异步后端与多数据源协作，项目中覆盖缓存、消息队列、文档存储、全文检索和前后端联调。",
        "文档意识较强，能维护 PRD、接口说明和测试脚本，适合 AI 应用开发、Python 后端开发、政企数字化系统开发等岗位。",
    ]:
        add_bullet(doc, item)

    add_section_heading(doc, "补充信息")
    add_bullet(doc, "求职方向：AI应用开发工程师、Python后端开发工程师、后端开发实习生。")
    add_bullet(doc, "可面试展开：FastAPI 异步接口设计、Redis Geo 派单、RabbitMQ 延迟/死信队列、Dify 工作流接入、ES 聚合统计。")

    doc.save(OUT)


if __name__ == "__main__":
    build()
