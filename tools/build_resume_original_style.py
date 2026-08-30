# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"C:\Users\yangshiqi\Desktop\2300130314杨诗琪_参考原版样式.docx"

FONT = "Microsoft YaHei"
TITLE = "333333"
BODY = "666666"
LIGHT_TEXT = "FFFFFF"
SECTION = "4F6F8F"
SECTION_DARK = "3E5D78"
SIDEBAR = "F4F6F8"
PHOTO = "D9DEE5"
BORDER = "D7DDE5"


def set_run(run, size=9.5, bold=False, color=BODY, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_run(p, text, size=9.5, bold=False, color=BODY):
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return run


def set_para(p, before=0, after=2, line=1.0, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    b = tc_pr.first_child_found_in("w:tcBorders")
    if b is None:
        b = OxmlElement("w:tcBorders")
        tc_pr.append(b)
    for edge in ("top", "left", "bottom", "right"):
        el = b.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            b.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def heading(cell, title):
    t = cell.add_table(rows=1, cols=1)
    table_width(t, 2000)
    c = t.cell(0, 0)
    shade(c, SECTION)
    borders(c, SECTION, "1")
    cell_margins(c, top=55, bottom=55, start=80, end=80)
    p = c.paragraphs[0]
    set_para(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, title, size=10.0, bold=True, color=LIGHT_TEXT)


def main_heading(cell, title):
    p = cell.add_paragraph()
    set_para(p, before=4, after=2)
    add_run(p, title, size=11.0, bold=True, color=TITLE)
    p_pr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), SECTION)
    bdr.append(bottom)
    p_pr.append(bdr)


def bullet(cell, text, size=8.6):
    p = cell.add_paragraph(style="List Bullet")
    set_para(p, after=1, line=1.02)
    add_run(p, text, size=size, color=BODY)


def info_line(cell, label, value):
    p = cell.add_paragraph()
    set_para(p, after=2, line=1.0)
    add_run(p, label, size=8.4, bold=True, color=TITLE)
    add_run(p, value, size=8.4, color=BODY)


def project(cell, title, date, stack, bullets):
    p = cell.add_paragraph()
    set_para(p, before=3, after=0, line=1.0)
    add_run(p, title, size=9.4, bold=True, color=TITLE)
    add_run(p, f"  {date}", size=8.2, color=SECTION_DARK)
    s = cell.add_paragraph()
    set_para(s, after=1, line=1.0)
    add_run(s, stack, size=7.9, color="777777")
    for item in bullets:
        bullet(cell, item, size=8.35)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(0.85)
    sec.bottom_margin = Cm(0.75)
    sec.left_margin = Cm(0.95)
    sec.right_margin = Cm(0.95)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 1.0

    bullet_style = styles["List Bullet"]
    bullet_style.font.name = FONT
    bullet_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    bullet_style.font.size = Pt(8.4)
    bullet_style.paragraph_format.left_indent = Cm(0.34)
    bullet_style.paragraph_format.first_line_indent = Cm(-0.16)
    bullet_style.paragraph_format.space_after = Pt(1)
    bullet_style.paragraph_format.line_spacing = 1.02

    layout = doc.add_table(rows=1, cols=2)
    layout.autofit = False
    table_width(layout, 9360)
    left = layout.cell(0, 0)
    right = layout.cell(0, 1)
    left.width = Cm(4.75)
    right.width = Cm(12.05)
    for c in (left, right):
        borders(c, "FFFFFF", "1")
        cell_margins(c, top=0, bottom=0, start=90, end=110)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    shade(left, SIDEBAR)
    shade(right, "FFFFFF")

    # Left sidebar: photo, identity, details.
    photo = left.add_table(rows=1, cols=1)
    table_width(photo, 1850)
    pc = photo.cell(0, 0)
    shade(pc, PHOTO)
    borders(pc, "C8D0DA", "5")
    cell_margins(pc, top=360, bottom=360, start=80, end=80)
    pp = pc.paragraphs[0]
    set_para(pp, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(pp, "照片", size=10.5, bold=True, color="777777")

    p = left.add_paragraph()
    set_para(p, before=8, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "杨诗琪", size=19.5, bold=True, color=TITLE)
    p = left.add_paragraph()
    set_para(p, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "AI应用开发 / Python后端开发", size=8.7, bold=True, color=SECTION_DARK)

    heading(left, "基本信息")
    info_line(left, "电话：", "130xxxxxxxx")
    info_line(left, "邮箱：", "xxxxxx@xxx.com")
    info_line(left, "居住地：", "XXXX")
    info_line(left, "毕业时间：", "2027届")

    heading(left, "教育背景")
    p = left.add_paragraph()
    set_para(p, before=3, after=1, line=1.0)
    add_run(p, "软件学院", size=8.8, bold=True, color=TITLE)
    p = left.add_paragraph()
    set_para(p, after=1)
    add_run(p, "计算机科学与技术 / 本科", size=8.3, color=BODY)
    p = left.add_paragraph()
    set_para(p, after=4)
    add_run(p, "2023.09 - 2027.06", size=8.3, color=BODY)

    heading(left, "专业技能")
    for item in [
        "Python、FastAPI、Pydantic、SQLAlchemy Async",
        "MySQL、Redis、MongoDB、Elasticsearch",
        "RabbitMQ、JWT、RESTful API、异步服务",
        "Dify 工作流、Prompt、AI 结构化解析",
        "Vue 3、Element Plus、ECharts、Git",
    ]:
        bullet(left, item, size=8.1)

    heading(left, "求职方向")
    for item in ["AI应用开发工程师", "Python后端开发工程师", "后端开发实习生"]:
        bullet(left, item, size=8.1)

    # Right main content.
    title = right.paragraphs[0]
    set_para(title, after=0)
    add_run(title, "个人简历", size=22, bold=True, color=SECTION_DARK)
    sub = right.add_paragraph()
    set_para(sub, after=5)
    add_run(sub, "FastAPI 异步后端 · Dify AI 工作流 · 多数据源系统设计", size=9.0, color="777777")

    main_heading(right, "专业总结")
    p = right.add_paragraph()
    set_para(p, after=3, line=1.08)
    add_run(
        p,
        "计算机科学与技术本科在读，主攻 Python 后端与 AI 应用工程化。具备 FastAPI 异步服务、"
        "SQLAlchemy/MySQL、Redis、MongoDB、Elasticsearch、RabbitMQ、Dify 工作流及 Vue 前端协作经验；"
        "能独立完成从需求分析、数据建模、接口开发到联调测试的完整项目闭环。",
        size=8.7,
        color=BODY,
    )

    main_heading(right, "项目经历")
    project(
        right,
        "城市公共设施智能报修与派单系统",
        "2026.06",
        "FastAPI / Vue 3 / MySQL / Redis / MongoDB / Elasticsearch / RabbitMQ / Dify",
        [
            "负责核心后端业务编排：实现市民报修、工单热缓存、图片元数据、AI解析日志、ES同步消息、超时派单检查等链路；以 MySQL 作为同步落地点，其余增强能力异步容错。",
            "设计四库分层数据模型：MySQL 承载事务数据，Redis 承载缓存、Geo空间计算和分布式锁，MongoDB 承载维修记录/附件/审计/AI日志，Elasticsearch 支撑全文检索与统计聚合。",
            "实现智能派单流程：基于 Redis Geo 筛选候选维修员，引入高德驾车距离修正与5分钟缓存，使用 Redis SETNX 防止并发双派，并通过 RabbitMQ 延迟消息处理无人接单兜底。",
            "参与三端页面联调：市民端报修/进度、维修员H5接单/完工/绩效、管理后台工单/调度/设施/结算/驾驶舱模块，配合 ECharts 展示运营指标。",
        ],
    )
    project(
        right,
        "Blog Log AI System 博客日志AI分析系统",
        "2026.05",
        "FastAPI / SQLAlchemy Async / MySQL / Dify API / Jinja2 / JavaScript",
        [
            "独立搭建博客管理与 AI 分析平台，完成注册登录、JWT 鉴权、博客 CRUD、分页查询、标题唯一性校验、Markdown 内容管理和用户数据隔离。",
            "封装 Dify 工作流调用，支持单篇博客摘要/关键词/难度/分类分析、全局主题聚类分析、个性化学习路线推荐，并将原始响应与结构化结果落库。",
            "建设统计分析接口：实现总博客数、月度新增、总字数、今日字数、发布/草稿分布、TOP博客、写作时段等指标；编写 Markdown 清洗函数，提升字数统计准确性。",
            "按 API-Service-CRUD-Model 分层组织代码，配套统一日志、异常处理、响应结构和测试脚本，降低接口扩展与问题定位成本。",
        ],
    )

    main_heading(right, "个人优势")
    for item in [
        "能把 AI 能力落到具体业务闭环中，关注结构化输出、失败降级、日志追踪和数据沉淀。",
        "熟悉异步后端与多数据源协作，项目覆盖缓存、消息队列、文档存储、全文检索和前后端联调。",
        "文档意识较强，能维护 PRD、接口说明和测试脚本，适合 AI 应用开发、Python 后端开发、政企数字化系统开发等岗位。",
    ]:
        bullet(right, item, size=8.35)

    main_heading(right, "可面试展开")
    p = right.add_paragraph()
    set_para(p, after=0, line=1.05)
    add_run(
        p,
        "FastAPI 异步接口设计、Redis Geo 派单、RabbitMQ 延迟/死信队列、Dify 工作流接入、ES 聚合统计。",
        size=8.45,
        color=BODY,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
